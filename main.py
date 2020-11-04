import docx                                                                                             # pip install python-docx
from nltk.stem import SnowballStemmer                                                                   # pip install nltk
from nltk.corpus import stopwords
from nltk import tokenize, FreqDist
from tika import parser                                                                                 # pip install tika
from langdetect import detect                                                                           # pip install langdetect
import re
import pandas as pd                                                                                     # pip install pandas
import os

fdist1 = ""
language = {
    "nl": "dutch",
    "en": "english",
    "es": "spanish",
    "fr": "french",
    "de": "german",
    "it": "italian"
    }

with open("keywords.txt", "r") as keywords:                                                             # Import file with keywords
    keywords = keywords.read().split()                                                                  # Convert keywords into list

directory_path = input('What is the path of the directory? ')                                           # Select file
directory = os.fsencode(directory_path)

report = docx.Document()                                                                                # Create report document
report.add_heading(f'Analysis {os.path.basename(directory_path)}', 0)                                   # Add title to report

for file in os.listdir(directory):
    document_path = os.path.join(directory, file).decode()
    document = parser.from_file(document_path)                                                          # Retrieve text from file
    document = document['content']
    content = re.sub(r'http\S+', " ", document)                                                         # Delete all links
    content = re.sub("[^a-zA-Z0-9|^-]", " ", content).lower()                                           # Delete all punctuation/upper case letters
    content_words = tokenize.word_tokenize(content)                                                     # Split words into list
    language_name = language[detect(content)]                                                           # Detect text language
    content_words_core = [w for w in content_words if w not in stopwords.words(language_name)]          # Delete adverbs
    content_words_core = " ".join(filter(lambda x: x in keywords, content_words_core)).split()          # Delete all words except for words in keywords
    stemmed_words = [SnowballStemmer(language_name).stem(word) for word in content_words_core]          # Group different forms of a word to a single item
    for words in stemmed_words:
        fdist1 = FreqDist(stemmed_words)                                                                # Count occurrence of words
    top_10_words = pd.DataFrame(fdist1.most_common(10), columns=['Word', 'Count'])                      # Put top 10 words in table
    top_10_words.loc[len(top_10_words), ['Word', 'Count']] = ['Total', top_10_words['Count'].sum()]     # Add totals-row to table
    top_10_words["Count"] = top_10_words["Count"].astype('int')                                         # Change datatype of column 'Count' to integer

    filename = os.fsdecode(file)
    title = report.add_heading(filename, level=1)                                                       # Add subtitle per document
    text_language = report.add_paragraph(f'Language: {language_name.capitalize()}')                     # Add language
    table = report.add_table(top_10_words.shape[0]+1, top_10_words.shape[1])                            # Add template table
    for j in range(top_10_words.shape[-1]):                                                             # Add headers to table
        table.cell(0, j).text = top_10_words.columns[j]
    for i in range(top_10_words.shape[0]):                                                              # Add data to table
        for j in range(top_10_words.shape[-1]):
            table.cell(i+1, j).text = str(top_10_words.values[i, j])
    table.style = 'Light Shading'                                                                       # Change style of table


report.save(f'{os.environ["USERPROFILE"]}/Desktop/report.docx')                                         # Save document to desktop
